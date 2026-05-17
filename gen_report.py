#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
生成 iBatis to JDBC 转换方案报告（Word 格式）
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime


def add_heading_with_style(doc, text, level):
    """添加标题"""
    heading = doc.add_heading(text, level=level)
    return heading


def add_paragraph_with_style(doc, text, bold=False, italic=False):
    """添加段落"""
    p = doc.add_paragraph(text)
    if bold or italic:
        for run in p.runs:
            if bold:
                run.bold = True
            if italic:
                run.italic = True
    return p


def add_table(doc, rows, cols):
    """添加表格"""
    return doc.add_table(rows=rows, cols=cols)


def generate_report():
    """生成报告"""
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(11)
    
    # ============ 标题页 ============
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('iBatis 到 JDBC 转换方案')
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_run.font.name = '宋体'
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run('技术方案与测试说明')
    subtitle_run.font.size = Pt(16)
    subtitle_run.font.name = '宋体'
    
    # 日期和版本
    doc.add_paragraph()
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_p.add_run(f'生成日期：{datetime.now().strftime("%Y年%m月%d日")}')
    date_run.font.size = Pt(11)
    
    version_p = doc.add_paragraph()
    version_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    version_run = version_p.add_run('版本：v1.2')
    version_run.font.size = Pt(11)
    
    # 分页符
    doc.add_page_break()
    
    # ============ 目录 ============
    add_heading_with_style(doc, '目录', 1)
    # 换行符号
    doc.add_paragraph()
    doc.add_paragraph('1. 项目概述')
    doc.add_paragraph('2. 转换方案设计')
    doc.add_paragraph('3. 核心功能说明')
    doc.add_paragraph('4. 系统架构')
    doc.add_paragraph('5. 实现细节')
    doc.add_paragraph('6. 测试验证')
    doc.add_paragraph('7. 适用场景与价值')
    doc.add_paragraph('8. 后续扩展规划')
    doc.add_paragraph('9. 技术栈')
    doc.add_paragraph('10. 总结')
    doc.add_paragraph('11. 落地计划（里程碑）')
    
    doc.add_page_break()
    
    # ============ 1. 项目概述 ============
    add_heading_with_style(doc, '1. 项目概述', 1)
    
    add_heading_with_style(doc, '1.1 背景与目标', 2)
    doc.add_paragraph(
        '目前系统中 iBatis 已无法继续使用，但项目中仍留下大量 SQL 定义在 XML 文件中。'
        '由于 iBatis 停止维护，直接升级到 MyBatis 或其他 ORM 框架会面临兼容性问题、成本高、风险大等障碍。'
        '不得不转向使用 JDBC 来替代，但 JDBC 本身无法识别 iBatis XML 文件。'
        '手工将每条 SQL 从 XML 转换为代码工作量巨大，且测试验证成本极高（堪称地狱级别）。'
        '本项目通过自动化方案，智能解析 iBatis XML，自动展开动态标签，生成可供 JDBC 直接执行的 SQL 与参数列表，'
        '极大降低迁移成本与测试压力。'
    )
    
    add_heading_with_style(doc, '1.2 核心价值', 2)
    doc.add_paragraph('降低迁移成本：批量转换 SQL，减少人工改写工作量', style='List Bullet')
    doc.add_paragraph('提升可验证性：通过测试与报告输出支持逐条核对', style='List Bullet')
    doc.add_paragraph('增强可读性：生成 Markdown、SQL 与 Excel 明细，便于审查', style='List Bullet')
    doc.add_paragraph('保持兼容性：覆盖常见 iBatis 动态 SQL 语法', style='List Bullet')
    
    # ============ 2. 转换方案设计 ============
    add_heading_with_style(doc, '2. 转换方案设计', 1)
    
    add_heading_with_style(doc, '2.1 总体思路', 2)
    doc.add_paragraph('转换流程分为两个阶段：')
    doc.add_paragraph('第一阶段：XML 解析与标签展开')
    doc.add_paragraph(
        '读取 iBatis SQLMap XML 文件，识别所有 SQL 语句定义，'
        '展开动态标签（dynamic、iterate等），根据示例参数完成参数内联。',
        style='List Bullet 2'
    )
    doc.add_paragraph('第二阶段：JDBC 标准化')
    doc.add_paragraph(
        '将内联后的 SQL 转换为 JDBC 标准格式：'
        '用 ? 占位符替换参数，维护有序的参数绑定列表。',
        style='List Bullet 2'
    )
    
    add_heading_with_style(doc, '2.2 核心转换规则', 2)
    
    table = add_table(doc, 6, 3)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'iBatis 占位符'
    hdr_cells[1].text = '转换规则'
    hdr_cells[2].text = '示例'
    
    rules = [
        ('#param#', '参数替换为 ? + 绑定值列表', '#id# → ? （id: 1001）'),
        ('$param$', '参数原样拼接 SQL（不带引号）', '$orderBy$ → id DESC'),
        ('dynamic 标签', '根据参数有效性展开或删除条件', '参数为空时跳过该子句'),
        ('iterate 标签', '遍历集合展开多个占位符', 'ids: [1, 2, 3] → IN (?, ?, ?)'),
        ('条件标签', '根据比较表达式决定是否渲染', 'isNotNull/isEqual/isGreaterThan 等'),
    ]
    
    for i, (marker, rule, example) in enumerate(rules, 1):
        cells = table.rows[i].cells
        cells[0].text = marker
        cells[1].text = rule
        cells[2].text = example

    add_heading_with_style(doc, '2.3 iBatis2 参数类型兼容策略（上线防错）', 2)
    doc.add_paragraph(
        '为避免迁移后线上参数类型错配导致数据库层异常，系统在转换入口增加 parameterClass 校验与规范化，'
        '策略以“与 iBatis2 行为对齐 + 应用层 fail-fast”为原则。'
    )

    table2 = add_table(doc, 7, 4)
    table2.style = 'Light Grid Accent 1'
    hdr2 = table2.rows[0].cells
    hdr2[0].text = '声明类型（parameterClass）'
    hdr2[1].text = '输入示例'
    hdr2[2].text = '处理结果'
    hdr2[3].text = '兼容性结论'

    compatibility_rules = [
        ('long', '"123"', '转为 123L', '与 iBatis2 常见行为一致（数值字符串可解析）'),
        ('long', '123 (Integer)', '转为 123L', '与 iBatis2 常见行为一致（数值类型可归一）'),
        ('long', '""', '抛异常（empty string cannot be converted to number）', '与 iBatis2 严格类型语义一致'),
        ('long', 'true', '抛异常（expected numeric value）', '与 iBatis2 严格类型语义一致'),
        ('long', 'null', '保留 null（不转 0）', '与 iBatis2 默认行为一致（未配置 nullValue）'),
        ('map', 'List / String', '抛异常（expected map-like object）', '与 iBatis2 参数契约一致'),
    ]

    for i, (declared_type, input_sample, result, conclusion) in enumerate(compatibility_rules, 1):
        cells = table2.rows[i].cells
        cells[0].text = declared_type
        cells[1].text = input_sample
        cells[2].text = result
        cells[3].text = conclusion
    
    # ============ 3. 核心功能说明 ============
    add_heading_with_style(doc, '3. 核心功能说明', 1)
    
    add_heading_with_style(doc, '3.1 支持的 iBatis 特性', 2)
    
    doc.add_paragraph('参数占位符：')
    doc.add_paragraph('#param# - 参数替换为 ?（JDBC prepared 模式）', style='List Bullet 2')
    doc.add_paragraph('$param$ - 参数原样拼接（SQL 片段内联模式）', style='List Bullet 2')
    doc.add_paragraph('多种参数类型自动推断：Map、String、Number、Bean、List、Array、Enum', style='List Bullet 2')
    
    doc.add_paragraph('条件标签：')
    doc.add_paragraph('dynamic、trim - 条件段落展开', style='List Bullet 2')
    doc.add_paragraph('isNotEmpty、isEmpty - 空值判断', style='List Bullet 2')
    doc.add_paragraph('isNull、isNotNull - NULL 判断', style='List Bullet 2')
    doc.add_paragraph('isEqual、isNotEqual - 等值比较', style='List Bullet 2')
    doc.add_paragraph('isGreaterThan、isLessThan - 大小比较', style='List Bullet 2')
    doc.add_paragraph('compareValue、compareProperty - 自定义比较', style='List Bullet 2')
    
    doc.add_paragraph('集合处理：')
    doc.add_paragraph('iterate - 遍历集合生成 IN 语句', style='List Bullet 2')
    doc.add_paragraph('Collection、Array、List 参数别名', style='List Bullet 2')
    
    doc.add_paragraph('其他特性：')
    doc.add_paragraph('include - 引用 SQL 片段', style='List Bullet 2')
    doc.add_paragraph('resultMap - 映射结果集', style='List Bullet 2')
    doc.add_paragraph('CDATA 字符清理 - 处理特殊字符', style='List Bullet 2')
    
    add_heading_with_style(doc, '3.2 参数类型支持', 2)
    doc.add_paragraph('Map - 命名参数', style='List Bullet')
    doc.add_paragraph('String / Number - 简单标量类型', style='List Bullet')
    doc.add_paragraph('JavaBean - 属性反射提取', style='List Bullet')
    doc.add_paragraph('List / Array / Collection - 集合参数', style='List Bullet')
    doc.add_paragraph('Enum - 枚举值比较', style='List Bullet')
    doc.add_paragraph('null - 特殊处理', style='List Bullet')
    
    # ============ 4. 系统架构 ============
    add_heading_with_style(doc, '4. 系统架构', 1)
    
    add_heading_with_style(doc, '4.1 核心类设计', 2)
    
    doc.add_paragraph('IbatisToJdbcConverter：主转换器')
    doc.add_paragraph('入口方法：convertPrepared(statementId, parameters)', style='List Bullet 2')
    doc.add_paragraph('扫描方法：loadSqlMapsFromClasspath(fileOrDirPaths)', style='List Bullet 2')
    doc.add_paragraph('扫描结果：getLoadedStatementInfos()/getLoadedStatementCount()', style='List Bullet 2')
    doc.add_paragraph('职责：转换 SQL、提取参数绑定', style='List Bullet 2')

    doc.add_paragraph('ConvertedSql：转换结果对象')
    doc.add_paragraph('核心字段：sql（? 占位符形式）、preparedBindings（参数列表）', style='List Bullet 2')
    doc.add_paragraph('常用方法：toPreviewSql()（生成可执行 SQL 预览）', style='List Bullet 2')

    doc.add_paragraph('JdbcExecutor：JDBC 执行器')
    doc.add_paragraph('职责：将转换结果按参数顺序绑定并执行', style='List Bullet 2')

    doc.add_paragraph('IbatisXmlSupport：XML 工具类')
    doc.add_paragraph('职责：标签识别、占位符解析、属性提取', style='List Bullet 2')
    
    add_heading_with_style(doc, '4.2 处理流程', 2)
    
    doc.add_paragraph('1. 扫描并加载 SQLMap 文件')
    doc.add_paragraph('通过 loadSqlMapsFromClasspath 扫描 classpath、目录或 jar 中的 *_sqlmap.xml 文件', style='List Bullet 2')
    doc.add_paragraph('解析 XML 并建立 statement/sql/resultMap 索引到新快照', style='List Bullet 2')
    doc.add_paragraph('原子切换生效，所有新请求立即使用新索引', style='List Bullet 2')
    
    doc.add_paragraph('2. 查询 statement 定义')
    doc.add_paragraph('根据 statementId 从 activeSnapshot 定位对应的 SQL 语句', style='List Bullet 2')
    
    doc.add_paragraph('3. 展开动态标签')
    doc.add_paragraph('递归处理 dynamic、iterate 等条件标签', style='List Bullet 2')
    
    doc.add_paragraph('4. 参数内联或占位')
    doc.add_paragraph('根据转换模式（内联/prepared）处理参数', style='List Bullet 2')
    
    doc.add_paragraph('5. SQL 正规化')
    doc.add_paragraph('清理多余空白、转换为标准 JDBC 格式', style='List Bullet 2')
    
    doc.add_paragraph('6. 返回转换结果')
    doc.add_paragraph('ConvertedSql 对象包含 SQL、参数、元数据和 preparedBindings（JDBC参数列表）', style='List Bullet 2')
    
    # ============ 5. 实现细节 ============
    add_heading_with_style(doc, '5. 实现细节', 1)
    
    add_heading_with_style(doc, '5.1 参数占位符处理', 2)
    doc.add_paragraph(
        '系统采用"双模式"设计：'
    )
    doc.add_paragraph(
        '预览模式（toPreviewSql()）：将 # 占位符替换为实际参数值（带类型转换），',
        style='List Bullet'
    )
    doc.add_paragraph(
        '便于人工审查和 SQL 预览',
        style='List Bullet 2'
    )
    doc.add_paragraph(
        'JDBC 模式（getPreparedBindings()）：将 # 占位符转为 ?，',
        style='List Bullet'
    )
    doc.add_paragraph(
        '维护有序的参数列表（preparedBindings），供 PreparedStatement 按顺序绑定',
        style='List Bullet 2'
    )
    
    add_heading_with_style(doc, '5.2 动态标签展开逻辑', 2)
    doc.add_paragraph(
        'dynamic 标签根据内层条件决定是否保留子内容：'
    )
    doc.add_paragraph('条件为真（非空、非零、条件成立）→ 保留内容', style='List Bullet')
    doc.add_paragraph('条件为假 → 删除内容', style='List Bullet')
    doc.add_paragraph('trim 标签用于移除多余的 AND/OR/逗号', style='List Bullet')
    
    add_heading_with_style(doc, '5.3 集合迭代处理', 2)
    doc.add_paragraph(
        'iterate 标签遍历集合，为每个元素生成一个占位符：'
    )
    doc.add_paragraph('输入：ids = [1, 2, 3]', style='List Bullet')
    doc.add_paragraph('输出：IN (?, ?, ?)，参数绑定：[1, 2, 3]', style='List Bullet')
    
    add_heading_with_style(doc, '5.4 XML热更新机制（双缓冲+原子切换）', 2)
    doc.add_paragraph(
        '为支持SQLMap的在线热更新，系统采用"双缓冲+原子切换"无锁设计，确保热更新过程中的高并发安全：'
    )
    doc.add_paragraph(
        '加载新XML：后台线程构建完整的新快照（statement索引、sql片段、resultMap等全量数据）',
        style='List Bullet'
    )
    doc.add_paragraph(
        '原子切换：新快照构建完成后，通过AtomicReference一次性原子切换指针，瞬间完成切换',
        style='List Bullet'
    )
    doc.add_paragraph(
        '读取无锁：所有请求直接读volatile activeSnapshot，完全无需加锁，避免热更新期间的性能下降',
        style='List Bullet'
    )
    doc.add_paragraph(
        '并发安全保证：',
        style='List Bullet'
    )
    doc.add_paragraph(
        '旧请求继续使用旧快照，新请求立即使用新快照，新旧快照不会交织',
        style='List Bullet 2'
    )
    doc.add_paragraph(
        '热更新时旧请求不受阻，新请求不等待，零延迟、零阻塞',
        style='List Bullet 2'
    )
    doc.add_paragraph(
        '旧快照自动垃圾回收，无需手动管理',
        style='List Bullet 2'
    )
    doc.add_paragraph(
        '运维优势：无需重启应用即可加载最新SQLMap文件，提升系统可用性和灵活性',
        style='List Bullet'
    )
    
    # ============ 6. 测试验证 ============
    add_heading_with_style(doc, '6. 测试验证', 1)
    
    add_heading_with_style(doc, '6.1 单元测试覆盖', 2)
    doc.add_paragraph('基础转换：SELECT/INSERT/UPDATE/DELETE 语句', style='List Bullet')
    doc.add_paragraph('参数类型：Map/String/Number/Bean/List/Array/Enum', style='List Bullet')
    doc.add_paragraph('条件标签：dynamic/trim/isNotEmpty/isNull/isGreaterThan 等', style='List Bullet')
    doc.add_paragraph('集合处理：iterate、集合参数别名', style='List Bullet')
    doc.add_paragraph('异常处理：statement 不存在、include 引用丢失等', style='List Bullet')
    
    add_heading_with_style(doc, '6.2 测试运行', 2)
    doc.add_paragraph('运行命令：mvn clean test', style='List Bullet')
    doc.add_paragraph('关键回归命令：mvn clean -q -Dtest=IbatisToJdbcConverterTest test', style='List Bullet')
    doc.add_paragraph('测试规模：IbatisToJdbcConverterTest（当前 19 个用例）', style='List Bullet')
    doc.add_paragraph('测试结果：通过（最近一次执行 TEST_EXIT:0）', style='List Bullet')

    add_heading_with_style(doc, '6.4 关键新增验证点（本轮）', 2)
    doc.add_paragraph('新增泛型查询能力：queryForObject/queryForList(Class<T>) 与 query(RowMapper<T>)', style='List Bullet')
    doc.add_paragraph('新增 iBatis2 参数类型对齐校验：在 convertPrepared 入口进行 parameterClass 验证', style='List Bullet')
    doc.add_paragraph('参数兼容综合测试：已合并为单一测试方法，覆盖 11 个场景并附注释', style='List Bullet')
    doc.add_paragraph('异常测试整合：statement/include/resultMap 缺失场景已合并并通过', style='List Bullet')
    
    add_heading_with_style(doc, '6.3 报告生成', 2)
    doc.add_paragraph('SQLMap 报告：自动读取 SQLMap，生成 Markdown 与 SQL 脚本', style='List Bullet')
    doc.add_paragraph('Excel 导出：从 Markdown 报告提取结构化字段并生成明细表', style='List Bullet')
    doc.add_paragraph('覆盖范围：以当前测试资源目录中的样本文件为准', style='List Bullet')
    
    doc.add_page_break()
    
    # ============ 7. 适用场景与价值 ============
    add_heading_with_style(doc, '7. 适用场景与价值', 1)
    
    add_heading_with_style(doc, '7.1 主要应用场景', 2)
    doc.add_paragraph('iBatis → JDBC 迁移：批量转换 SQL 语句', style='List Bullet')
    doc.add_paragraph('iBatis → Spring Data JPA 迁移：理解原始 SQL 结构', style='List Bullet')
    doc.add_paragraph('iBatis → 自研 ORM 框架：生成标准化 SQL', style='List Bullet')
    doc.add_paragraph('遗留系统升级：文档化和验证现有 SQL 逻辑', style='List Bullet')
    
    add_heading_with_style(doc, '7.2 业务价值', 2)
    doc.add_paragraph('加速迁移：自动化转换，显著减少手工改写工作量', style='List Bullet')
    doc.add_paragraph('降低风险：通过测试验证转换准确性', style='List Bullet')
    doc.add_paragraph('提高质量：确保参数顺序、SQL 逻辑一致', style='List Bullet')
    doc.add_paragraph('便于审查：生成可读性强的报告，便于人工检查', style='List Bullet')
    doc.add_paragraph('热更新支持：无需重启应用即可加载最新 SQLMap，大幅提升运维效率', style='List Bullet')
    doc.add_paragraph('高并发安全：采用双缓冲+原子切换，在高并发场景下无需加锁', style='List Bullet')
    doc.add_paragraph('综合最优：在兼容性、迁移成本与落地风险约束下，可视为当前阶段可能的最优解', style='List Bullet')
    
    # ============ 8. 后续扩展规划 ============
    add_heading_with_style(doc, '8. 后续扩展规划', 1)
    
    add_heading_with_style(doc, '8.1 功能扩展', 2)
    doc.add_paragraph('支持更多 iBatis 标签：placeholder、typeAlias 等', style='List Bullet')

    add_heading_with_style(doc, '8.3 风险边界与治理建议', 2)
    doc.add_paragraph('SQL 注入风险：$param$ 为原样拼接，必须在调用侧对来源与白名单做限制', style='List Bullet')
    doc.add_paragraph('类型契约风险：parameterClass 与实参不一致会 fail-fast，建议迁移期先做全量扫描', style='List Bullet')
    doc.add_paragraph('兼容性风险：个别历史 SQL 可能依赖隐式类型转换，需通过报告逐条核验', style='List Bullet')
    doc.add_paragraph('运维风险：热更新需纳入变更流程（灰度、回滚、审计）', style='List Bullet')
    
    add_heading_with_style(doc, '8.2 性能优化（未来方向）', 2)
    doc.add_paragraph('分片索引：对超大规模 SQLMap 进行分片优化', style='List Bullet')
    doc.add_paragraph('智能缓存失效：根据文件变动细粒度更新索引', style='List Bullet')
    
    # ============ 9. 技术栈 ============
    add_heading_with_style(doc, '9. 技术栈', 1)

    add_heading_with_style(doc, '9.1 依赖与第三方包说明', 2)
    doc.add_paragraph('生产代码依赖：仅使用 Java 8 标准库（java.* / javax.* / org.w3c.dom / org.xml.sax）', style='List Bullet')
    doc.add_paragraph('第三方运行时依赖：无（不依赖 Spring、MyBatis、Apache Commons 等外部库）', style='List Bullet')
    doc.add_paragraph('测试依赖：仅引入 JUnit 5（org.junit.jupiter），作用域为 test，不进入生产运行时', style='List Bullet')
    
    doc.add_paragraph('语言：Java 8', style='List Bullet')
    doc.add_paragraph('构建：Maven 3.9+', style='List Bullet')
    doc.add_paragraph('测试框架：JUnit 5', style='List Bullet')
    doc.add_paragraph('XML 解析：DOM / XPath', style='List Bullet')
    doc.add_paragraph('其他：正则表达式、反射、集合操作', style='List Bullet')
    
    # ============ 10. 总结 ============
    add_heading_with_style(doc, '10. 总结', 1)
    
    doc.add_paragraph(
        '本项目提供了可落地的 iBatis 到 JDBC 转换方案。'
        '通过自动化处理 XML 解析、动态标签展开与参数绑定，'
        '可在迁移过程中提升效率并降低人工改写风险。'
        '结合测试与报告，可支持逐条核对与持续迭代。'
    )
    
    doc.add_paragraph()
    doc.add_paragraph(
        '该方案适用于遗留系统现代化改造场景，可在保持业务连续性的前提下分阶段推进。'
        '在现有系统约束与迁移目标下，该方案可视为当前阶段可能的最优解。'
    )

    # ============ 11. 落地计划（里程碑） ============
    add_heading_with_style(doc, '11. 落地计划（里程碑）', 1)

    add_heading_with_style(doc, '11.1 里程碑划分', 2)
    doc.add_paragraph('M1【已完成】：已完成 SQLMap 清单扫描与高风险语句标注（$param$、复杂 dynamic）', style='List Bullet')
    doc.add_paragraph('M2【已完成】：已完成转换器接入与 parameterClass 类型错配治理', style='List Bullet')
    doc.add_paragraph('M3【未开始】：完成项目中 iBatis 调用的全量替换（切换到 JdbcExecutor/SpringJdbcExecutor）并完成联调', style='List Bullet')
    doc.add_paragraph('M4【未开始】：上线后监控与热更新治理，持续优化边角场景（所有 SQL 均为运行时自动转换，后续如有新增 XML 也可自动支持，无需手工静态替换）', style='List Bullet')

    add_heading_with_style(doc, '11.2 验收标准', 2)
    doc.add_paragraph('功能验收：关键业务 SQL 全量通过转换并可执行（所有 SQL 均在运行时自动转换，无需静态替换）；被融合项目的 iBatis 调用已全量替换为新方法，未替换项为 0', style='List Bullet')
    doc.add_paragraph('质量验收：核心测试通过率 100%，无阻断级回归问题', style='List Bullet')
    doc.add_paragraph('运行验收：系统稳定运行，无严重故障或高优先级问题', style='List Bullet')
    
    # 保存文档
    output_path = 'iBatis转JDBC转换方案报告.docx'
    doc.save(output_path)
    print(f'✅ 报告已生成：{output_path}')


if __name__ == '__main__':
    generate_report()
